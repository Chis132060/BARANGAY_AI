"""Create an honest, system-aligned copy of the Smart Barangay thesis DOCX."""

from pathlib import Path
import sys
import shutil

DOCX_TOOL_DIR = Path(r"C:\Users\donne\AppData\Local\Temp\barangay_ai_docx_tools")
sys.path.insert(0, str(DOCX_TOOL_DIR))

from docx import Document


SOURCE = Path(r"C:\Users\donne\Downloads\Copy of the final Barangay AI - IT32.docx")
REPO_OUTPUT = Path(r"C:\Users\donne\Videos\BARANGAY_AI\docs\BARANGAY_AI_CAPSTONE_THESIS_REVISED.docx")
DOWNLOAD_OUTPUT = Path(r"C:\Users\donne\Downloads\Smart_Barangay_Capstone_Thesis_Revised.docx")


def set_paragraph(document, index, text):
    """Replace paragraph text while retaining its paragraph style."""
    paragraph = document.paragraphs[index]
    paragraph.text = text


def clear_paragraph(document, index):
    set_paragraph(document, index, "")


def revise(input_path: Path, output_path: Path):
    if not input_path.exists():
        raise FileNotFoundError(f"Thesis DOCX not found: {input_path}")

    document = Document(str(input_path))

    # Chapter 1 framing: describe the current web/PWA implementation and avoid
    # presenting target capabilities as completed automation.
    set_paragraph(document, 30, "This chapter introduces the Smart Barangay Portal, an AI-assisted web and Progressive Web Application (PWA) system developed for Barangay Tandang Sora, Butuan City. It outlines the background of the study, review of related literature, statement of the problem and objectives, theoretical, conceptual, and research paradigm, scope and limitations, and definition of terms. This revised manuscript distinguishes verified current functionality from planned enhancements.")
    set_paragraph(document, 70, "The integration of web and mobile access is particularly important for barangay services because many residents rely on smartphones as their primary device for communication and information access. BarangayConnect is accessible via mobile browsers but does not include a dedicated mobile application, and the authors explicitly recommend a separate mobile app to improve accessibility and user experience (Enaño et al., 2025). SMART BARANGAY responds to this need through a web portal for staff and a mobile-friendly Progressive Web Application (PWA) for residents, using a shared backend for the features currently implemented.")
    set_paragraph(document, 72, "Integrated web and mobile-friendly platforms can use API-based communication to share data between front-end applications and a back-end database. In SMART BARANGAY, the current build supports shared authentication, request storage, status tracking, and administrative verification flows through Supabase. Complete synchronization, real-time service updates, and all end-to-end document workflows remain subject to deployment and integration testing.")
    set_paragraph(document, 79, "Intelligent government systems that combine AI with structured workflows can automate routine inquiries, guide citizens through procedures, and escalate complex cases to human personnel. This supports SMART BARANGAY's goal of assisting with repetitive questions and routing service requests while keeping staff responsible for decisions that require human judgment. In the current build, the chatbot provides knowledge-based answers and a basic request-intake flow; broader workflow automation is treated as an implementation target until verified end to end.")
    set_paragraph(document, 96, "However, LLM-powered chatbots in government settings require strict control to ensure that responses remain accurate, appropriate, and aligned with official policies. SMART BARANGAY is configured to use a barangay knowledge base, RAG-capable API services, and session identifiers. The final deployment must ensure that only approved sources are retrieved, unsupported questions receive a safe fallback, and privacy and retention controls are tested before compliance is claimed.")

    # Related literature: change the project-specific workflow and memory claims
    # from completed automation to the verified partial implementation.
    set_paragraph(document, 106, "Digital Request Management and Limited Workflow Automation")
    set_paragraph(document, 110, "SMART BARANGAY applies these ideas through a chatbot-assisted intake flow. When a resident describes a need, such as requesting a barangay clearance, the current interface can identify a related service and display a basic request card. This is an assisted intake feature rather than a complete service-specific dynamic form.")
    clear_paragraph(document, 111)
    set_paragraph(document, 112, "After submission, the request is stored for status tracking and administrative review when the resident record is correctly linked. Required fields, document uploads, automatic approval decisions, filled-template generation, payment processing, and ready-for-pickup delivery are not yet a single verified end-to-end workflow.")
    set_paragraph(document, 116, "In SMART BARANGAY, the chatbot client creates a session identifier and sends it with chat requests to support conversational context. The current persistence path does not yet prove complete database-linked session memory for every message or a fully enforced multi-request session. Therefore, session context is described as a design capability under verification, not as permanent cross-session memory.")
    clear_paragraph(document, 117)
    set_paragraph(document, 118, "The Philippine Data Privacy Act of 2012 (Republic Act No. 10173) requires personal data to be handled securely and only for legitimate, authorized purposes. The system must enforce user/session isolation, a documented retention period, access controls, and a tested deletion procedure. These controls should be reported as implemented only after deployment verification.")
    set_paragraph(document, 125, "In the context of the Smart Barangay Portal, the DeLone and McLean model informs how system quality, information quality, service quality, user satisfaction, and net benefits will be evaluated across the staff web portal, resident PWA, and chatbot. The current build provides a basis for testing responsiveness, knowledge-answer quality, request tracking, and usability. Claims about notification reliability, processing efficiency, reduced in-person visits, and net benefits must be supported by actual UAT and performance evidence rather than projected outcomes.")
    set_paragraph(document, 139, "The Smart Barangay Portal's AI-assisted chatbot is grounded in Conversational AI and NLP theory through its API and RAG-oriented implementation. It accepts natural-language resident queries and maps them to informational responses or a basic service-intake flow. The client sends a session identifier to support conversational context, but complete persistent session memory, multi-request grouping, and retention/deletion enforcement require additional implementation and testing. The current system should therefore be described as context-assisted rather than as a fully implemented session-memory architecture.")

    # Conceptual framework and objectives.
    set_paragraph(document, 154, "The process component represents the current core operations of the system. It includes portal access and user identification, natural-language processing of resident queries, retrieval of configured knowledge, basic request-intake display, storage of submitted records, staff verification or request review where supported, and presentation of announcements and notifications. Service-specific dynamic forms, payment handling, filled-template generation, and automatic delivery are proposed extensions. Privacy controls must include explicit session ownership, retention, and deletion verification before the Data Privacy Act claim is stated as fully implemented.")
    set_paragraph(document, 155, "The output component refers to the results produced by the interaction between the resident and the system, as well as the administrative actions that follow. Current outputs include chatbot guidance, stored service requests for administrative review, registration verification states, request status displays, announcements, and in-app notification records where configured.")
    clear_paragraph(document, 156)
    set_paragraph(document, 157, "Automated report generation, generated official templates, payment notices, and external push/SMS delivery should be treated as target outputs unless the corresponding implementation and test evidence are included. The expected broader outcomes, such as fewer in-person visits and improved efficiency, remain evaluation hypotheses rather than guaranteed results.")

    set_paragraph(document, 168, "The main objective of this study is to develop SMART BARANGAY, an AI-assisted web and Progressive Web Application system that supports the administrative and constituent services of Barangay Tandang Sora, Butuan City.")
    set_paragraph(document, 170, "To achieve this objective, the study pursues the following specific objectives:")
    set_paragraph(document, 171, "To design and develop a staff web portal and resident-friendly Progressive Web Application with shared authentication and database access for the features implemented in the current build.")
    set_paragraph(document, 172, "To integrate a large language model (LLM) and retrieval-oriented knowledge services to assist with resident inquiries and identify possible service-request intents, subject to source validation and safe fallback behavior.")
    set_paragraph(document, 173, "To implement configured administrative services covering barangay knowledge content, announcements, resident verification, and monitoring of supported service requests. Service-specific documents such as clearances, certificates of indigency, proof of income, and barangay IDs should be claimed only when their fields, validation, workflow, and output are implemented and tested.")
    set_paragraph(document, 174, "To implement an AI-assisted chatbot that provides continuous access to configured barangay information and guides residents to the appropriate service or human review when the system cannot safely answer.")
    set_paragraph(document, 175, "To support a traceable digital request workflow from resident intake to staff review and status tracking. Full automation of document generation, payment collection, pickup notification, and final delivery is outside the currently verified scope unless completed and evaluated before submission.")
    set_paragraph(document, 182, "The system may lessen the burden of repetitive inquiries and manual recordkeeping by assisting with information dissemination, resident verification, request intake, and status monitoring. Barangay personnel remain responsible for reviewing submissions, approving or rejecting requests, issuing official documents, and performing tasks that require administrative judgment.")
    set_paragraph(document, 190, "Senior citizens, working residents, and other users with limited time or mobility may benefit from remote access to barangay information and supported services. The system's access controls and session handling are intended to protect user information, but retention, deletion, and notification privacy controls must be verified in the deployed environment.")
    set_paragraph(document, 193, "barangay-level digital service delivery and future workflow automation.")
    set_paragraph(document, 195, "This study focuses on the design, development, deployment, and evaluation of the Smart Barangay Portal for Barangay Tandang Sora, Butuan City. The system is composed of a web portal intended for barangay officials and administrators and a resident-facing Progressive Web Application (PWA), rather than a separately compiled native Android or iOS application.")
    set_paragraph(document, 196, "The scope of the study includes uploading and managing barangay rules, policies, and service guidelines; posting announcements and updates; resident registration with staff verification where the migration is deployed; receiving and monitoring supported service requests; providing AI-assisted responses to resident inquiries; and enabling residents to submit basic requests digitally. The study also covers the RAG-oriented chatbot route and request status tracking. Service-specific dynamic forms, attachments, filled templates, payments, external push/SMS delivery, automated report exports, and complete session deletion are limitations or future enhancements unless separately verified.")

    # Definitions must describe the actual implementation boundary.
    set_paragraph(document, 201, "Artificial Intelligence refers to the language-model and retrieval-assisted technology used by the Smart Barangay Portal to interpret resident inquiries and provide guided information; it does not independently approve official requests or replace barangay staff decisions.")
    set_paragraph(document, 210, "Notification refers to a stored or displayed in-app message used to inform users about configured announcements or request/verification updates. External push, SMS, and automatic status-triggered delivery require separate implementation and deployment verification.")
    set_paragraph(document, 217, "Session-Based Memory refers to the temporary conversational context associated with an active chatbot session. In this study, the intended behavior is limited, user-owned, and subject to explicit retention and deletion controls; it should not be interpreted as permanent memory across sessions.")
    set_paragraph(document, 222, "Workflow Automation refers to system-supported routing, storage, status tracking, or notification steps that reduce manual repetition. In the current implementation, automation is partial; official review, approval, document issuance, payment decisions, and delivery remain staff-controlled unless a specific workflow is demonstrated end to end.")

    # System description and privacy statements.
    set_paragraph(document, 240, "The Web Portal functions as a digital office for the barangay. Authorized staff can manage configured knowledge content, post announcements and community updates, verify resident registrations where the approval migration is deployed, and review supported service requests. Records are stored in Supabase for the implemented modules, while full operational coverage and production controls remain subject to deployment testing.")
    set_paragraph(document, 241, "The resident-facing PWA features an AI-assisted chatbot that uses configured barangay information and a RAG-capable API route when available. For supported intents, the interface can guide a resident and show a basic request-intake card. It does not yet provide a verified dynamic form for every document type, automatic filled-template generation, payment collection, or external pickup/delivery notification.")
    clear_paragraph(document, 242)
    clear_paragraph(document, 243)
    clear_paragraph(document, 244)
    set_paragraph(document, 245, "Resident privacy is a required design objective, not a completed claim by default. The system must enforce authenticated ownership, row-level access, session isolation, a documented retention period, and a tested deletion process. The final thesis should state compliance with Republic Act No. 10173 only for controls that are implemented, deployed, and supported by verification evidence.")
    set_paragraph(document, 246, "By integrating an AI-assisted chatbot, a digital records system, resident verification, and a supported request-status workflow, the Smart Barangay Portal aims to reduce unnecessary in-person inquiries and improve staff visibility. Staff review and official document issuance remain part of the process, while broader automation is a future enhancement unless demonstrated.")

    # Architecture and data-flow descriptions.
    set_paragraph(document, 284, "For resident-users, the portal directs the user to the AI-assisted chatbot, where an inquiry may be entered in Bisaya, Filipino, or English. The RAG-capable route or local knowledge fallback provides guidance when a matching topic is available. The current interface may display a basic request-intake card for supported intents, after which a request is stored for status tracking and possible administrative review. It should not be described as a complete dynamic-form, automatic-template, payment, or deletion workflow until those components are verified.")
    set_paragraph(document, 289, "For supported administrative tasks, staff review resident verification records, announcements, and service requests through the web portal. Approval or rejection remains a staff decision. In-app notification records and realtime verification updates are available where configured; automatic ready-for-pickup notifications, payment notices, push delivery, and SMS delivery require explicit status-triggered implementations and deployment testing.")
    set_paragraph(document, 290, "Resident and administrator data can support dashboard monitoring. A fully automated report-generation and export function is not claimed in this revised manuscript unless it is implemented, reproducible, and reconciled against database records.")
    set_paragraph(document, 302, "The Data Flow Diagram should be interpreted as the target and current-supported flow: residents submit inquiries or supported requests through the PWA, the chatbot/API provides guidance, records are stored in the database, and staff review or verification actions update the applicable records. Knowledge indexing, notifications, report generation, payment handling, and document generation are separate capabilities and must be labeled according to their actual deployment status.")
    set_paragraph(document, 304, "This flow supports digital service access, record organization, resident verification, and request monitoring. It does not by itself prove completed automatic document issuance, external notification delivery, payment processing, or report exports.")
    set_paragraph(document, 311, "The Use Case Diagram shows the intended interactions between residents, barangay staff, and the AI-assisted system. Residents can register, ask configured knowledge questions, submit supported requests, view announcements, and track statuses. Staff can verify residents, manage configured content, and review requests. Automated inquiry assistance and routing are supported, while official approval, document issuance, payment decisions, and complete report generation remain staff-controlled or future functions unless tested end to end.")
    set_paragraph(document, 319, "The ERD shows how the Smart Barangay Portal organizes barangay, resident, user, service-request, document, announcement, notification, and audit data. The current schema supports these relationships in configured environments, but the application must use the residents primary key consistently when creating requests. Chat-session persistence and document/embedding links must also be verified against the deployed schema before they are presented as complete working flows.")
    set_paragraph(document, 321, "This structure provides a foundation for record tracking and knowledge retrieval. Reliable multi-request sessions, generated documents, notification events, payment metadata, and report outputs require additional tables, services, or tested workflows where they are not yet present.")

    # Methodology: distinguish development activities from completed evaluation.
    set_paragraph(document, 251, "Implementation - This phase involved coding the current FastAPI-based API routes, RAG-oriented chatbot services, Supabase data layer, resident PWA, and administrative web portal according to the approved design. Features that remain partial, such as service-specific forms, template generation, payment handling, external notifications, and reporting exports, are recorded as limitations rather than treated as completed outputs.")
    clear_paragraph(document, 252)
    set_paragraph(document, 253, "Integration and Testing - This phase covers functional, security, and privacy testing of the implemented modules. Only test cases that were actually executed, including their date, environment, fixtures, and results, should be reported as findings.")
    clear_paragraph(document, 254)
    set_paragraph(document, 255, "Unverified target workflows must remain marked as pending validation.")
    set_paragraph(document, 256, "Deployment - This phase refers to configuring and releasing the verified build to its intended environment. A feature should be labeled deployed only after its migration, environment variables, external services, access controls, and end-to-end behavior have been checked in the target environment.")
    set_paragraph(document, 257, "Operation and Maintenance - This phase covers monitoring and issue resolution after deployment. If production monitoring or stakeholder use has not yet occurred, it should be reported as planned work rather than completed operation.")
    set_paragraph(document, 373, "Phase 2 System Development. Following the needs assessment, the team developed the current AI-assisted chatbot, administrative dashboard, resident PWA, request-intake flow, resident verification workflow, and database/API integrations. The implementation uses the documented TypeScript/Next.js, FastAPI, Supabase, and AI service components. A complete automated reporting module, automatic document-template generation, payment workflow, and external notification delivery are not claimed unless the corresponding code and deployment evidence are included.")
    set_paragraph(document, 376, "For UAT, participants should be given access only to the workflows available in the tested build, such as registration, resident verification, asking a knowledge question, submitting a supported request, and viewing status. Tasks such as generating a sample report or completing an automatic document workflow must be removed from the completed-results narrative unless they were actually available and tested. The manuscript should include the actual participant count and collection dates.")
    set_paragraph(document, 377, "Phase 4 System Performance Benchmarking. Performance evaluation should report only executed measurements, such as chatbot response time, request submission time, dashboard load time, or notification delivery time, with the test environment and sample size stated. A 200-query intent-accuracy test or report-generation benchmark must not be presented as a completed result without raw queries, scoring criteria, timings, and computed results.")

    # Add a compact revision note at the end of the document without disturbing
    # the original references and existing page layout.
    document.add_heading("Revision Note - System Alignment", level=1)
    document.add_paragraph("This revision was prepared from the current code audit at repository revision 0d23a14. The thesis now distinguishes implemented capabilities from partial or planned capabilities. The revised research emphasizes AI-assisted information access, resident verification, supported digital request intake, staff-controlled processing, and status tracking. Full automation of document generation, payment, pickup/delivery notifications, report exports, and complete session deletion requires additional implementation and evidence.")
    document.add_paragraph("Original source preserved: Copy of the final Barangay AI - IT32.docx. This revised copy is saved under a new filename; the source file was not overwritten.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main():
    revise(SOURCE, REPO_OUTPUT)
    shutil.copy2(REPO_OUTPUT, DOWNLOAD_OUTPUT)
    print(f"Revised DOCX: {DOWNLOAD_OUTPUT}")
    print(f"Repository copy: {REPO_OUTPUT}")


if __name__ == "__main__":
    main()
